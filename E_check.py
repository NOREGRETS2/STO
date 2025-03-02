import processing_CFG_in_txt
from docx import Document
import re
#Проверка названия файла


#Массив копия со всеми возможными названиями
array_parametr_filename_t = processing_CFG_in_txt.array_parametr_filename

def parameters_with_fio(array_parametr_filename_t):
    array_parametr_with_fio = []
    for string in array_parametr_filename_t:
        if string.endswith(" fio\n"):
            array_parametr_with_fio.append(string)
    return array_parametr_with_fio

def parametrs_with_not_fio(array_parametr_filename_t):
    array_parametr_with_not_fio = []
    for string in array_parametr_filename_t:
        if string.endswith(" not_fio\n"):
            array_parametr_with_not_fio.append(string)
    return array_parametr_with_not_fio

def remove_fio(element):
    return element.replace(" fio\n", "")

def remove_n_fio(element):
    return element.replace(" not_fio\n", "")

array_parametr_with_fio = parameters_with_fio(array_parametr_filename_t)
array_parametr_with_not_fio = parametrs_with_not_fio(array_parametr_filename_t)

array_parametr_with_fio = [remove_fio(element) for element in array_parametr_with_fio]
array_parametr_with_not_fio = [remove_n_fio(element) for element in array_parametr_with_not_fio]

def extract_name(file_name):
    match = re.search(r'.*\/(.*\.docx)$', file_name)
    if match:
        return match.group(1)
    else:
        return None


def trim_name(file_name):
    last_index = file_name.rfind(' ')
    if file_name != -1:
        return file_name[:last_index]
    else:
        return file_name

def error_3_1(file_name, array_parametr_with_fio):
    a = 0
    error31 = ''
    for item in array_parametr_with_fio:
        if item == file_name:
            a = 1
    if a != 1:
        error31 = 'Ошибка пункт СТО 3.1'
        return error31


def error_4_1_6(file_name):
    document = Document(file_name)
    error416 = ''
    found_page_number = False

    for section in document.sections:
        header = section.header
        footer = section.footer
        if any('PAGE' in paragraph.text for paragraph in header.paragraphs):
            found_page_number = True
            break
        if any('PAGE' in paragraph.text for paragraph in footer.paragraphs):
            found_page_number = True
            break

    if not found_page_number:
        error416 = 'Ошибка пункт СТО 4.1.6'

    return error416

file_name = processing_CFG_in_txt.name_and_path_file
error416 = error_4_1_6(file_name)
file_name = extract_name(file_name)
file_name = trim_name(file_name)
error31 = error_3_1(file_name, array_parametr_with_fio)
#print(file_name)
#print(array_parametr_with_fio)

#print(error31,error416)
#print(array_parametr_with_not_fio)

