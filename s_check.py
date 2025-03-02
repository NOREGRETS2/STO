import processing_CFG_in_txt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

file = processing_CFG_in_txt.name_and_path_file


def podtiraem_mas(mas):
    clean = [i.replace('\n', '') for i in mas]
    return clean


def find_word(file, ar):
    doc = Document(file)
    for i, paragraph in enumerate(doc.paragraphs):
        if (ar[0] in paragraph.text) and (i >= 33):
            return i


def check_intro(file, ar):
    doc = Document(file)
    for i, paragraph in enumerate(doc.paragraphs):
        if (ar[0] in paragraph.text) and (paragraph.alignment != ar[5]) and len(paragraph.text) < 2:
            mas_s[0].append(i)
            mas_s[1].append('Несоответствие пункту СТО 4.5.5')
            break
        else:
            mas_s[0].append(0)
            mas_s[1].append('Несоответствие пункту СТО 4.5.5')
            break

ar = podtiraem_mas(processing_CFG_in_txt.array_alexandr)
mas_s = [[],[]]
check_intro(file, ar)