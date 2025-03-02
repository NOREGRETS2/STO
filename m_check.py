import processing_CFG_in_txt
from docx import Document
from docx.enum.text import WD_LINE_SPACING

# Получаем имя и путь к файлу
name_and_path_file = processing_CFG_in_txt.name_and_path_file
maxim = processing_CFG_in_txt.array_maxim

# Инициализируем переменные для минимального и максимального значений шрифта
min_score = ""
max_score = ""

# Список для хранения ошибок
m_errors = [[], []]

# Устанавливаем минимальные и максимальные значения шрифта на основе массива maxim
if maxim[0] == "0\n":
    min_score = int(12)
if maxim[1] == "1\n":
    max_score = int(14)



# Функция для проверки размера шрифта
def check_font_size(paragraphs, min_score, max_score):
    size_errors = []  # Список для хранения ошибок размера шрифта
    for idx, paragraph in enumerate(paragraphs):  # Перебираем абзацы
        for run in paragraph.runs:  # Перебираем части текста в абзаце
            try:
                # Проверяем, находится ли размер шрифта в допустимых пределах
                if run.font.size is not None and not (min_score <= run.font.size.pt <= max_score):
                    size_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})  # Добавляем ошибку
                    break  # Выходим из цикла, если ошибка найдена
            except AttributeError:
                # Обрабатываем случай, если атрибут size отсутствует
                size_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})
                break
    return size_errors  # Возвращаем список ошибок

# Функция для проверки межстрочного интервала
def check_line_spacing(paragraphs, min_score, max_score):
    spacing_errors = []  # Список для хранения ошибок межстрочного интервала
    for idx, paragraph in enumerate(paragraphs):  # Перебираем абзацы
        for run in paragraph.runs:  # Перебираем части текста в абзаце
            try:
                # Проверяем, соответствует ли межстрочный интервал заданным условиям
                if run.font.size is not None:
                    if run.font.size.pt == min_score and \
                            paragraph.paragraph_format.line_spacing_rule != WD_LINE_SPACING.SINGLE or \
                            run.font.size.pt == max_score and \
                            paragraph.paragraph_format.line_spacing_rule != WD_LINE_SPACING.ONE_POINT_FIVE:
                        spacing_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})  # Добавляем ошибку
                        break  # Выходим из цикла, если ошибка найдена
            except AttributeError:
                # Обрабатываем случай, если атрибут size отсутствует
                spacing_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})
                break  # Выходим из цикла
    return spacing_errors

# Основная функция
def process_file(name_and_path_file: str) -> None:
    try:
        doc = Document(name_and_path_file)  # Открываем документ
        font_size_errors = check_font_size(doc.paragraphs, min_score, max_score)  # Проверяем размер шрифта
        if font_size_errors:  # Если есть ошибки
            for error in font_size_errors:
                m_errors[0].append("Ошибка пункт СТО 4.1.5")  # Добавляем описание ошибки
                m_errors[1].append(error["paragraph_text"])  # Добавляем текст абзаца с ошибкой

        line_spacing_errors = check_line_spacing(doc.paragraphs, min_score, max_score)  # Проверяем межстрочный интервал
        if line_spacing_errors:  # Если есть ошибки
            for error in line_spacing_errors:
                m_errors[0].append("Ошибка пункт СТО 4.1.5")  # Добавляем описание ошибки
                m_errors[1].append(error["paragraph_text"])  # Добавляем текст абзаца с ошибкой

    except Exception as e:
        print(f"Произошла ошибка при обработке файла {e} ")  # Обрабатываем исключения и выводим сообщение об ошибке

process_file(name_and_path_file)