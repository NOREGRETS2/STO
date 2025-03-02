from docx import Document
import processing_CFG_in_txt
from docx.shared import Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH


kostya = processing_CFG_in_txt.array_kostya
file_name = processing_CFG_in_txt.name_and_path_file
massiv_k = [[], []]

viravnivanie = ""
if kostya[0] == "0\n":
    viravnivanie = WD_ALIGN_PARAGRAPH.JUSTIFY
#print(viravnivanie)

markeri = ""
if kostya[1] == "1\n":
    markeri = "*"
#print(markeri)

embed = ""
if kostya[2] == "2\n":
    port = float(30)
#print(port)

morph = ""
if kostya[3] == "3\n":
    zenith = float(20)
#print(zenith)

pulse = ""
if kostya[4] == "4\n":
    starboard = float(10)
#print(starboard)

cipher = ""
if kostya[5] == "5\n":
    scarlet = float(12.5)
#print(scarlet)

#Функция проверки на наличие в тексте маркеров
def check_list_styles(paragraphs, markeri):
    marker_errors = []
    for idx, paragraph in enumerate(paragraphs):
        if markeri in paragraph.text:
            marker_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})
    return marker_errors

# Функция проверки выравнивание абзацев по ширине
def check_alignment(paragraphs, viravnivanie):
    alignment_errors = []
    for idx, paragraph in enumerate(paragraphs):
        if not paragraph.style.name.startswith("Heading") and paragraph.alignment != viravnivanie and ("введение" not in paragraph.text.lower() and "заключение" not in paragraph.text.lower()):
            alignment_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})
    return alignment_errors



# Конвертация inches в миллиметры (мм)
def inches_to_mm(inches):
    return inches * 25.4

# Конвертация EMU в миллиметры (мм)
def emu_to_mm(emu):
    return round(float(Emu(emu).mm), 2)

def process_file(file_name: str) -> None:

    try:
        doc = Document(file_name)
        section = doc.sections[0]
        left_margin = Mm(30)
        top_margin = Mm(20)
        right_margin = Mm(10)
        bottom_margin = Mm(20)

        margin_ok = abs(inches_to_mm(section.left_margin.inches) - left_margin.mm) < 0.1 and \
                    abs(inches_to_mm(section.top_margin.inches) - top_margin.mm) < 0.1 and \
                    abs(inches_to_mm(section.right_margin.inches) - right_margin.mm) < 0.1 and \
                    abs(inches_to_mm(section.bottom_margin.inches) - bottom_margin.mm) < 0.1


        alignment_errors = check_alignment(doc.paragraphs,viravnivanie)
        if alignment_errors:
            for error in alignment_errors:
                massiv_k[0].append("Ошибка пункт СТО 4.1.5")
                massiv_k[1].append(error["paragraph_text"])

        markeri_errors = check_list_styles(doc.paragraphs,markeri)
        if markeri_errors:
            for error in markeri_errors:
                massiv_k[0].append("Ошибка пункт СТО 4.7")
                massiv_k[1].append(error["paragraph_text"])

    except Exception as e:
        print(f"Произошла ошибка при обработке файла: {e}")

process_file(file_name)





