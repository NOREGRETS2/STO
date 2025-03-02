import processing_CFG_in_txt
from docx import Document
from docx.shared import Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH

andrey = processing_CFG_in_txt.array_andrey
file_name = processing_CFG_in_txt.name_and_path_file

errors = [[], []]


glyph = ""
if andrey[0] == "0\n":
    glyph = "Times New Roman"
#print(glyph)

chroma = ""
if andrey[1] == "1\n":
    chroma = "000000"
#print(chroma)

embed = ""
if andrey[2] == "2\n":
    port = float(30)
#print(port)

morph = ""
if andrey[3] == "3\n":
    zenith = float(20)
#print(zenith)

pulse = ""
if andrey[4] == "4\n":
    starboard = float(10)
#print(starboard)

cipher = ""
if andrey[5] == "5\n":
    scarlet = float(12.5)
#print(scarlet)

shift = ""
if andrey[6] == "6\n":
    shift = "Arial"
#print(shift)

datamorpher = ""
if andrey[7] == "7\n":
    translate = int(12)
#print(translate)

translatron = ""
if andrey[8] == "8\n":
    metamorphosis = int(13)
#print(metamorphosis)

metamorphicator = ""
if andrey[9] == "9\n":
    convert = int(14)
#print(convert)

datashifter = ""
if andrey[10] == "10\n":
    magicvert = int(15)
#print(magicvert)

# Проверка выравнивания абзацев
def check_alignment(paragraphs, expected_alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    alignment_errors = []
    for idx, paragraph in enumerate(paragraphs):
        if paragraph.alignment != expected_alignment:
            alignment_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})
    return alignment_errors

# Конвертация inches в миллиметры
def inches_to_mm(inches):
    return inches * 25.4

# Конвертация EMU в миллиметры
def emu_to_mm(emu):
    return round(float(Emu(emu).mm), 2)

# Проверка отступа красной строки
def check_indent(paragraphs, expected_indent):
    indent_errors = []
    for idx, para in enumerate(paragraphs):
        if para.paragraph_format.first_line_indent is not None and not \
                abs(emu_to_mm(para.paragraph_format.first_line_indent) - expected_indent) < 0.1:
            indent_errors.append(
                {"paragraph_index": idx + 1, "paragraph_text": para.text}
            )
    return indent_errors

# Проверка шрифта
def check_font(paragraphs, glyph):
    font_errors = []
    for idx, paragraph in enumerate(paragraphs):
        for run in paragraph.runs:
            if run.font.name != glyph:
                font_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})
                break
    return font_errors

# Проверка цвета шрифта
def check_font_color(paragraphs, chroma):
    color_errors = []
    for idx, paragraph in enumerate(paragraphs):
        for run in paragraph.runs:
            if run.font.color.rgb is not None:
                r, g, b = run.font.color.rgb
                color_hex = f"{r:02x}{g:02x}{b:02x}"
                if color_hex != chroma:
                    color_errors.append({"paragraph_index": idx + 1, "paragraph_text": paragraph.text})
                    break
    return color_errors

# Основная функция
def process_file(file_name: str) -> None:
    try:
        doc = Document(file_name)

        section = doc.sections[0]
        left_margin = Mm(port)
        top_margin = Mm(zenith)
        right_margin = Mm(starboard)
        bottom_margin = Mm(zenith)
        first_line_indent = Mm(scarlet)

        margin_ok = abs(inches_to_mm(section.left_margin.inches) - left_margin.mm) < 0.1 and \
                    abs(inches_to_mm(section.top_margin.inches) - top_margin.mm) < 0.1 and \
                    abs(inches_to_mm(section.right_margin.inches) - right_margin.mm) < 0.1 and \
                    abs(inches_to_mm(section.bottom_margin.inches) - bottom_margin.mm) < 0.1

        if not margin_ok:
            errors[0].append("Ошибка пункт СТО 4.1.3")
            errors[1].append("Полe документа")

        alignment_errors = check_alignment(doc.paragraphs)
        if alignment_errors:
            for error in alignment_errors:
                errors[0].append("Ошибка пункт СТО 4.1.5")
                errors[1].append(error["paragraph_text"])

        font_errors = check_font(doc.paragraphs, glyph)
        if font_errors:
            for error in font_errors:
                errors[0].append("Ошибка пункт СТО 4.1.5")
                errors[1].append(error["paragraph_text"])

        font_color_errors = check_font_color(doc.paragraphs, chroma)
        if font_color_errors:
            for error in font_color_errors:
                errors[0].append("Ошибка пункт СТО 4.1.5")
                errors[1].append(error["paragraph_text"])

        indent_errors = check_indent(doc.paragraphs, first_line_indent.mm)
        if indent_errors:
            for error in indent_errors:
                errors[0].append("Ошибка пункт СТО 4.1.5")
                errors[1].append(error["paragraph_text"])

    except Exception as e:
        print(f"Произошла ошибка при обработке файла {e}")



process_file(file_name)
